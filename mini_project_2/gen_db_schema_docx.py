from docx import Document

doc = Document()
doc.add_heading('สรุปโครงสร้างตารางฐานข้อมูล (yumpooma.db)', 0)

tables = [
    {
        'name': 'bills',
        'columns': [
            ('id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT'),
            ('table_no', 'INTEGER', 'เลขโต๊ะ'),
            ('bill_time', 'TEXT', 'วันที่-เวลาออกบิล'),
            ('items', 'TEXT', 'JSON รายการอาหารในบิล'),
            ('total', 'REAL', 'ราคารวม'),
        ]
    },
    {
        'name': 'sales',
        'columns': [
            ('id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT'),
            ('sale_date', 'TEXT', 'วันที่ขาย (YYYY-MM-DD)'),
            ('menu_id', 'INTEGER', 'รหัสเมนู'),
            ('quantity', 'INTEGER', 'จำนวนที่ขาย'),
            ('total_price', 'REAL', 'ราคารวมของเมนูนี้'),
        ]
    },
    {
        'name': 'menus',
        'columns': [
            ('id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT'),
            ('name', 'TEXT', 'ชื่อเมนู'),
            ('price', 'REAL', 'ราคา'),
            ('description', 'TEXT', 'รายละเอียด (nullable)'),
            ('image', 'TEXT', 'ชื่อไฟล์รูปภาพ (nullable)'),
            ('category', 'TEXT', 'หมวดหมู่เมนู'),
        ]
    },
    {
        'name': 'reservations',
        'columns': [
            ('id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT'),
            ('name', 'TEXT', 'ชื่อลูกค้า'),
            ('phone', 'TEXT', 'เบอร์โทร'),
            ('date', 'TEXT', 'วันที่จอง'),
            ('time', 'TEXT', 'เวลาจอง'),
            ('table_no', 'INTEGER', 'เลขโต๊ะ'),
            ('people', 'INTEGER', 'จำนวนคน'),
            ('status', 'TEXT', 'สถานะ (nullable)'),
        ]
    },
    {
        'name': 'admin',
        'columns': [
            ('id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT'),
            ('username', 'TEXT', 'ชื่อผู้ใช้'),
            ('password', 'TEXT', 'รหัสผ่าน'),
        ]
    },
]

for t in tables:
    doc.add_heading(f"ตาราง {t['name']}", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ชื่อคอลัมน์'
    hdr_cells[1].text = 'ประเภทข้อมูล'
    hdr_cells[2].text = 'รายละเอียด'
    for col in t['columns']:
        row_cells = table.add_row().cells
        row_cells[0].text = col[0]
        row_cells[1].text = col[1]
        row_cells[2].text = col[2]
    doc.add_paragraph('')

doc.save('db_schema.docx')
